#!/usr/bin/env python3
"""Generate test fixture files for the renaming-file-personal-style tests."""

import os

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")


def create_fixtures():
    os.makedirs(FIXTURES_DIR, exist_ok=True)

    # 1. Plain text file
    with open(os.path.join(FIXTURES_DIR, "sample.txt"), "w", encoding="utf-8") as f:
        f.write(
            "March 15, 2025\n\n"
            "Dear Vincent Bortone,\n\n"
            "Your Chase Bank statement is enclosed.\n\n"
            "Sincerely,\nChase Bank\n"
        )

    # 2. Text-based PDF via pymupdf
    import fitz

    pdf_path = os.path.join(FIXTURES_DIR, "sample.pdf")
    doc = fitz.open()
    page = doc.new_page()
    text = (
        "Statement Date: February 10, 2025\n\n"
        "Account Holder: Jennifer Bortone\n"
        "Bank of America\n\n"
        "Checking Account Summary\n"
        "Beginning Balance: $5,000.00\n"
        "Ending Balance: $4,750.00\n"
    )
    page.insert_text((72, 72), text, fontsize=12)
    doc.save(pdf_path)
    doc.close()

    # 3. DOCX file
    from docx import Document

    docx_path = os.path.join(FIXTURES_DIR, "sample.docx")
    document = Document()
    document.add_paragraph("Date: January 5, 2025")
    document.add_paragraph("")
    document.add_paragraph("To: Vincent and Jennifer Bortone")
    document.add_paragraph("From: Internal Revenue Service")
    document.add_paragraph("")
    document.add_paragraph("RE: 2024 Federal Tax Return Acknowledgment")
    document.add_paragraph("")
    document.add_paragraph(
        "This letter confirms receipt of your 2024 federal tax return."
    )
    document.save(docx_path)

    # 4. Large text file (for truncation test)
    large_path = os.path.join(FIXTURES_DIR, "large.txt")
    with open(large_path, "w", encoding="utf-8") as f:
        # Write more than 10,000 chars
        f.write("A" * 15000)

    # 5. CSV file
    csv_path = os.path.join(FIXTURES_DIR, "sample.csv")
    with open(csv_path, "w", encoding="utf-8") as f:
        f.write("date,description,amount\n")
        f.write("2025-03-01,Electric bill,$150.00\n")
        f.write("2025-03-15,Gas bill,$75.00\n")

    # 6. PNG image with text (for OCR test)
    from PIL import Image, ImageDraw, ImageFont

    img_path = os.path.join(FIXTURES_DIR, "sample.png")
    img = Image.new("RGB", (600, 200), color="white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 24)
    except OSError:
        font = ImageFont.load_default()
    draw.text((20, 20), "Invoice Date: April 1, 2025", fill="black", font=font)
    draw.text((20, 60), "To: Margaret Bortone", fill="black", font=font)
    draw.text((20, 100), "From: Acme Insurance", fill="black", font=font)
    img.save(img_path)

    print(f"Fixtures created in {FIXTURES_DIR}")


if __name__ == "__main__":
    create_fixtures()
